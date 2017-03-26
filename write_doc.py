#Virus:(W1f!@c$w)

from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.shared import *
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

import docx
import mid_db

def write_doc(ftitle):

    readerc,readernc=mid_db.content_creation(ftitle)
    #Function to split string of ips
    def split_ip(st,n):
        st=st.split("\n")
        st1='\n'.join(st[:n/2])
        st2='\n'.join(st[n/2:])
        return "\n"+st1,"\n"+st2


    #Function to merge both contents without duplicates
    def merge():
        final=[]
        for i in csv_reader:
            for j in xml_reader:
                #To add ips if duplicate is found (using title or CVE)
                if i[0]==j[0] or i[6]==j[6]:
                    i[3]+="|"+j[3]
                    xml_reader.remove(j)
            final.append(i)
        return final+xml_reader

    #Combined contents from the function
    #reader=merge()

    #Function for assigning suitable values for High,Medium,Low
    def assign_val(i):
        if i[4]=="High" or i[4]=="Critical":
            return 3
        elif i[4]=="Medium":
            return 2
        elif i[4]=="Low":
            return 1

    #To sort the contents based on Risk Rating
    readerc=sorted(readerc,key=assign_val,reverse=True)
    readernc=sorted(readernc,key=assign_val,reverse=True)

    #Function to add hyperlink to links in the Section- See Also
    def add_hyperlink(paragraph, url):
        text=url

        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id, )
        hyperlink.set(qn('w:history'), '1')

        # Create a w:r element
        new_run = OxmlElement('w:r')

        # Create a new w:rPr element
        rPr = OxmlElement('w:rPr')

        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'Hyperlink')

        # Join all the xml elements together add add the required text to the w:r element
        rPr.append(rStyle)
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        # Create a new Run object and add the hyperlink into it
        r = paragraph.add_run()
        r._r.append(hyperlink)

        # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
        r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        r.font.underline = True
        return r

    #Function to set table column width
    def set_column_width(column, width):
        for cell in column.cells:
            cell.width = width

    #Open from template document which contains all styles required for this document
    document = docx.Document("template.docx")

    for reader in (readerc,readernc):
        if reader==readerc:
            document.add_paragraph("Vulnerabilities with CVE:")
        if reader==readernc:
            document.add_paragraph("Vulnerabilities without CVE:")
        rno=1
        for row in reader:
            #Create bullet-in points for each IP:port
            st=""
            n=0
            ip=list(set(row[3].split("|"))) #Split it into individual ips from the string based on the delimiter |

            for i in ip:
                leng=len(i)
                if i[leng-2:]==":0":
                    i=i[:leng-2]
                if n!=0:
                    st+=u'\u2022'+i+"\n" #ASCII Value for bullet-in symbol
                n += 1

            #Add table heading
            table_head = document.add_table(1,1)
            table_head.autofit = False
            table_head.columns[0].width = Cm(15)
            table_head.rows[0].cells[0].paragraphs[0].add_run(str(rno)+"."+row[0]).bold=True

            #Construct primary table
            table=document.add_table(rows=6,cols=2)
            table.style='LightGrid-Accent6'
            table.autofit=False
            table.alignment=WD_TABLE_ALIGNMENT.CENTER

            #Setting column widths
            set_column_width(table.columns[0], Cm(4))
            set_column_width(table.columns[1], Cm(11))
            #set_column_width(table.columns[2], Cm(1.5))

            #Setting headings in column 1 for each row
            table.rows[0].cells[0].paragraphs[0].add_run("Description").bold=False
            #q=table.rows[0].cells[1].merge(table.rows[0].cells[2])
            table.rows[1].cells[0].text="\nAssessment Type\n"
            table.rows[2].cells[0].text="Affected IP's and services"
            table.rows[3].cells[0].text="\nRisk Rating\n"
            table.rows[4].cells[0].text="\nImpact\n"
            table.rows[5].cells[0].text="Solution"

            #Inserting the values from the csv file into each table
            table.rows[0].cells[1].text=row[1].replace("\n"," ")#Description
            table.rows[1].cells[1].text=row[2]#Assesment Type
            if n>6:
                p=table.rows[2].cells[1].add_table(1,2)
                p.rows[0].cells[0].text,p.rows[0].cells[1].text=split_ip(st,n)
            else:
                table.rows[2].cells[1].text="\n"+st#Affected IP's and services with bullet-ins

            #To assign Critical to High and to insert Risk rating and to color the cell with appropriate color
            if row[4]!="Critical":
                if row[4]=="High":
                    table.rows[3].cells[1].paragraphs[0].add_run('\n'+row[4]+'\n')
                    table.rows[3].cells[1].paragraphs[0].style="RedS"
                if row[4]=="Medium":
                    table.rows[3].cells[1].paragraphs[0].add_run('\n'+row[4]+'\n')
                    table.rows[3].cells[1].paragraphs[0].style="YellowS"
                if row[4]=="Low":
                    table.rows[3].cells[1].paragraphs[0].add_run('\n'+row[4]+'\n')
                    table.rows[3].cells[1].paragraphs[0].style="GreenS"
            else:
                table.rows[3].cells[1].paragraphs[0].add_run('\n'+"High\n")
                table.rows[3].cells[1].paragraphs[0].style="RedS"


            table.rows[5].cells[1].text=row[5].replace("\n"," ")+"\n"#Add solution
            if row[5]=="":
                table.rows[5].cells[1].text="None"
            if row[6]!="":#Contents of xml parsing doesn't have See Also
                if row[6]:
                    table.rows[5].cells[1].paragraphs[0].add_run('\nSee also:\n').bold=True#Add see also to the solution

                    for i in row[6].split("\n"):

                        table.rows[5].cells[1].paragraphs[0].add_run("\n")
                        add_hyperlink(table.rows[5].cells[1].paragraphs[0], i)#Add hyperlinked links below see also


            #Add spacing at end of each table
            document.add_paragraph('\n')
            rno+=1


    #Saving the document in demo.docx
    document.save(ftitle+'.docx')